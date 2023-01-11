# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.


from docx import Document
from docx.shared import Inches
document = Document()

x = 0
while x < 5:
    x += 1
    document.add_heading('Разработка сайтов', 0)
    p = document.add_paragraph('Различной сложности, быстро и недорого \n')
    p.add_run('Телефон: 87477360201\n').bold = True
    p.add_run('www.iskakovstudio.kz').bold = True
    document.add_picture('1.jpg', width=Inches(2.25))
    document.add_paragraph('\n\n')
document.add_page_break()
document.save('визитка.docx')



import openpyxl # Это модуль Python для чтения и записи Excel 2010 xlsx/xlsm/xltx/xltm файлов
from openpyxl.drawing.image import Image
from openpyxl.styles import Font
import os
import datetime
import time
try:
    book = openpyxl.Workbook() # Создание пустого документа
    # book = openpyxl.load_workbook('sample.xlsx') # Открытие существующего документа
    sheet = book.active # Активация документа
    x = 0
    while x < 5:
        x += 1
        sheet['A'+str(x * 5)] = 'Разработка сайтов'
        sheet['A' + str(1 + (x * 5))] = 'Различной сложности, быстро и недорого'
        sheet['A' + str(2 + (x * 5))] = 'Телефон: 87477360201'
        sheet['A' + str(3 + (x * 5))] = 'www.iskakovstudio.kz'
        cell = sheet.cell(row=(x * 5), column=1)
        cell.font = Font(name='Calibri', size=20, bold=True) # Установка фонта ячейки
        sheet.add_image(Image('2.jpg'), 'F'+str(x * 5))


    my_file = "sample.xlsx" # Имя файла
    book.save(my_file) # Сохранение файла на диск
    os.startfile(my_file) # Запуск файла (открытие MS Excel с созданным документом)
except Exception as a: # Обработка ошибок
    print("Error!")
    print(a)